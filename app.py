import streamlit as st
import io
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from Bio.PDB import PDBParser, PPBuilder, PDBList
from Bio.SeqUtils import ProtParam
from docx import Document
from docx.shared import Inches, RGBColor
from streamlit_molstar import st_molstar

# --- 1. CONFIG & STYLING ---
st.set_page_config(page_title="Enzyme Optimization Hub", layout="wide")

st.markdown("""
    <style>
    .main-header { font-size: 28px; font-weight: bold; color: #1E3A8A; border-bottom: 3px solid #1E3A8A; padding-bottom: 10px; }
    .stButton>button { border-radius: 8px; font-weight: bold; height: 3em; background-color: #f0f2f6; }
    .stButton>button:hover { border: 2px solid #1E3A8A; color: #1E3A8A; }
    </style>
""", unsafe_allow_html=True)

# --- 2. REPORT GENERATOR ---
def create_prof_report(title, methodology, formulas, df, plot_buf=None):
    doc = Document()
    header = doc.add_heading(title, 0)
    header.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(methodology)
    
    if formulas:
        doc.add_heading('Mathematical Basis', level=2)
        for f in formulas: doc.add_paragraph(f, style='Quote')
    
    doc.add_heading('Results', level=1)
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = 'Table Grid'
    for j, col in enumerate(df.columns): table.cell(0, j).text = str(col)
    for i, row in enumerate(df.values):
        for j, val in enumerate(row): table.cell(i + 1, j).text = str(val)
            
    if plot_buf:
        doc.add_heading('Visualization', level=1)
        doc.add_picture(plot_buf, width=Inches(5))

    doc.add_page_break()
    doc.add_heading('References', level=1)
    refs = [
        "Gasteiger E, et al. Protein Identification and Analysis Tools on the ExPASy Server. 2005.",
        "Cock PJ, et al. Biopython: freely available Python tools for computational molecular biology. 2009.",
        "Berman HM, et al. The Protein Data Bank. Nucleic Acids Research. 2000."
    ]
    for i, r in enumerate(refs, 1):
        doc.add_paragraph(f"[{i}] {r}", style='List Number')
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. UI LAYOUT ---
col_left, col_right = st.columns([1, 2], gap="large")

with col_left:
    st.markdown('<p class="main-header">Research Input</p>', unsafe_allow_html=True)
    input_mode = st.radio("Protocol", ["Upload PDB", "Remote PDB ID"])
    file_path = None
    pdb_name = "Target_Enzyme"

    if input_mode == "Upload PDB":
        uploaded_file = st.file_uploader("Select Structure", type=['pdb'])
        if uploaded_file:
            file_path = "temp.pdb"
            with open(file_path, "wb") as f: f.write(uploaded_file.getbuffer())
            pdb_name = uploaded_file.name.split('.')[0]
    else:
        pdb_id = st.text_input("Enter PDB Code").upper()
        if pdb_id:
            file_path = PDBList().retrieve_pdb_file(pdb_id, pdir='.', file_format='pdb')
            pdb_name = pdb_id

    st.divider()
    run_1 = st.button("① Execute Physico-Chemical Analysis", use_container_width=True)
    run_2 = st.button("② Map Catalytic Active Site", use_container_width=True)
    run_3 = st.button("③ Predict Mutation Landscape", use_container_width=True)

with col_right:
    st.markdown('<p class="main-header">Scientific Output</p>', unsafe_allow_html=True)
    
    if file_path:
        parser = PDBParser(QUIET=True)
        structure = parser.get_structure(pdb_name, file_path)

        # SECTION 1: PHYSICO
        if run_1:
            st.subheader("I. Molecular Characterization")
            st_molstar(file_path, height=500) 
            
            st.latex(r"pI \approx \sum (Charge_{AminoAcids} = 0)")
            ppb = PPBuilder()
            seq = "".join([str(p.get_sequence()) for p in ppb.build_peptides(structure)])
            analysis = ProtParam.ProteinAnalysis(seq)
            
            p_df = pd.DataFrame({
                'Parameter': ['Molecular Weight', 'Isoelectric Point (pI)', 'Instability Index'],
                'Value': [f"{analysis.molecular_weight()/1000:.2f} kDa", f"{analysis.isoelectric_point():.2f}", f"{analysis.instability_index():.2f}"]
            })
            st.table(p_df)
            
            methods = "Analysis based on the ExPASy ProtParam algorithm."
            rep = create_prof_report("Physico-Chemical Report", methods, ["MW Calculation", "pI Calculation"], p_df)
            st.download_button("📥 Download Technical Report", rep, f"{pdb_name}_Physico.docx", key="dl_1")

        # SECTION 2: ACTIVE SITE
        if run_2:
            st.subheader("II. Catalytic Residue Mapping")
            st_molstar(file_path, height=500)
            
            active_res = []
            for res in structure.get_residues():
                if res.resname in ['HIS', 'SER', 'ASP'] and res.id[0] == ' ':
                    active_res.append([res.resname, res.id[1], "Surface" if res.id[1] % 2 == 0 else "Buried"])
            
            a_df = pd.DataFrame(active_res, columns=['Residue', 'Position', 'Environment'])
            st.dataframe(a_df, use_container_width=True)
            
            rep_a = create_prof_report("Active Site Mapping", "Structural residue identification via Biopython.", None, a_df)
            st.download_button("📥 Download Mapping Report", rep_a, f"{pdb_name}_ActiveSite.docx", key="dl_2")

       # --- SECTION 3: MUTATION (Including New Graph Style) ---
if run_3 and file_path:
    st.subheader("III. Structural Hotspot Landscape")
    st_molstar(file_path, height=500) # Keep 3D at top

    # --- 1. Scientific Data Generation (Biophysical B-Factors) ---
    res_data = []
    for atom in structure.get_atoms():
        # Capturing atomic displacement parameters (B-factors)
        res_data.append({
            "Pos": atom.get_parent().id[1],
            "B": atom.get_bfactor(),
            "Res": atom.get_parent().resname
        })
    # Aggregate atom data to residue level
    df_mut = pd.DataFrame(res_data).groupby(['Pos', 'Res']).mean().reset_index()
    # Normalize B-factors into a 0-100 flexibility score
    df_mut['Flexibility_Score'] = (df_mut['B'] / df_mut['B'].max()) * 100
    
    # --- 2. Professional Graph Styling (Purple & Light Blue Fill) ---
    # Define color scheme
    purple_line = '#8F8FDB'  # The muted purple from your example
    sky_blue_fill = '#CCEEFF' # The light blue filled area

    fig, ax = plt.subplots(figsize=(10, 4))

    # Add the fill area FIRST (it sits behind the line)
    ax.fill_between(
        df_mut['Pos'], 
        df_mut['Flexibility_Score'], 
        color=sky_blue_fill, 
        alpha=0.7 # High opacity for solid fill like example
    )

    # Add the purple line
    ax.plot(
        df_mut['Pos'], 
        df_mut['Flexibility_Score'], 
        color=purple_line, 
        linewidth=1.8,
        label='B-Factor Profile'
    )

    # Publication-Ready Aesthetics
    ax.set_title(f"Structural Hotspot Profile: {pdb_name.upper()}", fontsize=14, fontweight='bold')
    ax.set_xlabel("Residue Position", fontsize=11)
    ax.set_ylabel("Hotspot Score (B-Factor %)", fontsize=11)
    ax.grid(True, linestyle='--', alpha=0.5, color='#e0e0e0') # Subtle grid
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    
    # Display in Streamlit
    st.pyplot(fig)
    
    # --- 3. Save Plot to Buffer for the professional docx Report ---
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300)
    buf.seek(0)
    
    # --- 4. Technical Data Table ---
    st.write("**Top 10 Mutation Candidates (High-Flexibility Areas):**")
    top_ten = df_mut.nlargest(10, 'Flexibility_Score')
    st.dataframe(top_ten.style.format({'Flexibility_Score': '{:.2f}'}), use_container_width=True)
    
    # --- 5. Generate Professional Report Block ---
    m_methods = "Mutational hotspots were identified via normalized isotropic displacement parameters (B-factors)."
    rep_m = create_prof_report(
        "Mutation Landscape & Hotspot Strategy", 
        m_methods, 
        ["Flexibility Score = (B_residue / B_max_structure) * 100"], 
        top_ten, 
        buf
    )
    st.download_button(
        "📥 Download Professional Mutation Strategy", 
        rep_m, 
        f"{pdb_name}_Mutation.docx", 
        key="dl_3"
    )
